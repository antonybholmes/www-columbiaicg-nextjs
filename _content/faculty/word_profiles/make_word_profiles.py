# -*- coding: utf-8 -*-
"""
Created on Fri Jul 31 12:10:51 2020

@author: nt0ny
"""


import frontmatter
import os
from bs4 import BeautifulSoup
import urllib
import collections
import json
import time
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches

for file in os.listdir('../'):
    if '.md' not in file or 'generic' in file:
        continue
    
    post = frontmatter.load(os.path.join('../', file))
    
    person = frontmatter.load('../../people/{}.md'.format(post['id']))
    
    with open('../../selected-publications/{}.json'.format(post['id'])) as f:
        publications = json.load(f)
    
    print(file)
    
    
    document = Document()
    
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    p = document.add_paragraph('{} {}'.format(person['firstName'], person['lastName']).upper())
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    font = run.font
    font.name = 'Arial'
    font.bold = True
    font.size = Pt(11)
    
    p = document.add_paragraph('Overview')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.add_break()
    font = run.font
    font.name = 'Arial'
    font.bold = True
    font.size = Pt(11)
    
    p = document.add_paragraph('Email: {}'.format(person['email']))
    run = p.runs[0]
    run.add_break()
    font = run.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    p = document.add_paragraph('Academic Appointments')
    run = p.runs[0]
    font = run.font
    font.name = 'Arial'
    font.bold = True
    font.size = Pt(11)
    
    # appointments
    
    for titles in person['titles']:
        if 'admin' not in titles:
            for title in titles.split(';'):
                p = document.add_paragraph(title, style='List Bullet')
                run = p.runs[0]
                font = run.font
                font.name = 'Arial'
                font.size = Pt(11)
    
    admin = False
    
    for titles in person['titles']:
        if 'admin' in titles:
            admin = True
            break
    
    if admin:
        p = document.add_paragraph('Administrative Appointments')
        run = p.runs[0]
        font = run.font
        font.name = 'Arial'
        font.bold = True
        font.size = Pt(11)
        
        for titles in person['titles']:
            if 'admin' in titles:
                titles = titles.replace('admin::', '')
                for title in titles.split(';'):
                    p = document.add_paragraph(title, style='List Bullet')
                    run = p.runs[0]
                    font = run.font
                    font.name = 'Arial'
                    font.size = Pt(11)
        
    #p = document.add_paragraph('Administrative Titles')
    run = p.runs[0]
    run.add_break()
    #font = run.font
    #font.name = 'Arial'
    #font.bold = True
    #font.size = Pt(11)
    
    p = document.add_paragraph('Research')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    font = run.font
    font.name = 'Arial'
    font.bold = True
    font.size = Pt(11)
    
    content = post.content
    
    # remove tags
    
    content = re.sub(r'<.+?>', '', content)
    
    p = document.add_paragraph(content)
    run = p.runs[0]
    #run.add_break()
    font = run.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    p = document.add_paragraph('Selected Publications')
    run = p.runs[0]
    font = run.font
    font.name = 'Arial'
    font.bold = True
    font.size = Pt(11)
    
    # order by date
    pub_order = collections.defaultdict(lambda: 
                                        collections.defaultdict(lambda: 
                                                                collections.defaultdict(lambda: collections.defaultdict(object))))
                                                                
    
    for pub in publications['publications']:
        pub_order[pub['year']][pub['month']][pub['day']][pub['title']] = pub
    
    
    for year in reversed(sorted(pub_order)):
        for month in reversed(sorted(pub_order[year])):
            for day in reversed(sorted(pub_order[year][month])):
                for title in sorted(pub_order[year][month][day]):
                    pub = pub_order[year][month][day][title]
                    p = document.add_paragraph(pub['title'], style='List Number')
                    run = p.runs[0]
                    run.add_break()
                    font = run.font
                    font.bold = True
                    font.name = 'Arial'
                    font.size = Pt(11)
                    run = p.add_run(pub['authors'])
                    run.add_break()
                    font = run.font
                    font.name = 'Arial'
                    font.size = Pt(11)
                    run = p.add_run('{}. {}.'.format(pub['journal'], pub['year']))
                    run.add_break()
                    font = run.font
                    font.name = 'Arial'
                    font.size = Pt(11)
                    
                    if pub['pmid'] != '' or pub['doi'] != '':
                        run = p.add_run('PMID: {}, DOI: {}'.format(pub['pmid'], pub['doi']))
                        run.add_break()
                        font = run.font
                        font.name = 'Arial'
                        font.size = Pt(11)
        
        # change margin
    for section in document.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    document.save('CUIMC_{}_VPS_Template_April2021.docx'.format(person['lastName'].replace('-', '')))
    
    #break
    